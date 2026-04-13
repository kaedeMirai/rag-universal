from __future__ import annotations

import contextlib
import io
import json
import re
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from xml.etree import ElementTree as ET

import pandas as pd

CONTROL_CHARACTERS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")
WHITESPACE_RE = re.compile(r"[ \t]+")
MULTI_NEWLINE_RE = re.compile(r"\n{3,}")


@dataclass(frozen=True)
class ExtractedSegment:
    text: str
    page_start: int | None = None
    page_end: int | None = None


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    segments: list[ExtractedSegment]


@dataclass(frozen=True)
class ChunkedSegment:
    chunk_index: int
    text: str
    page_start: int | None = None
    page_end: int | None = None
    source_locator: str = ""


def clean_extracted_text(text: str) -> str:
    if not text:
        return ""

    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = CONTROL_CHARACTERS_RE.sub("", cleaned)
    cleaned = WHITESPACE_RE.sub(" ", cleaned)
    cleaned = "\n".join(line.strip() for line in cleaned.splitlines())
    cleaned = MULTI_NEWLINE_RE.sub("\n\n", cleaned)
    return cleaned.strip()


def normalize_text_for_retrieval(text: str) -> str:
    return clean_extracted_text(text)


def format_source_locator(page_start: int | None, page_end: int | None) -> str:
    if not page_start:
        return ""
    if page_end and page_end != page_start:
        return f"стр. {page_start}-{page_end}"
    return f"стр. {page_start}"


def serialize_segments(segments: list[ExtractedSegment]) -> str:
    if not segments:
        return ""
    return json.dumps([asdict(segment) for segment in segments], ensure_ascii=False)


def deserialize_segments(raw_value: str) -> list[ExtractedSegment]:
    if not raw_value.strip():
        return []

    try:
        payload = json.loads(raw_value)
    except json.JSONDecodeError:
        return []

    output: list[ExtractedSegment] = []
    if not isinstance(payload, list):
        return output

    for item in payload:
        if not isinstance(item, dict):
            continue
        text = clean_extracted_text(str(item.get("text", "") or ""))
        if not text:
            continue
        page_start = _coerce_optional_int(item.get("page_start"))
        page_end = _coerce_optional_int(item.get("page_end"))
        output.append(
            ExtractedSegment(
                text=text,
                page_start=page_start,
                page_end=page_end,
            )
        )
    return output


def _coerce_optional_int(value: object) -> int | None:
    if value in (None, "", 0, "0"):
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _build_document(segments: list[ExtractedSegment]) -> ExtractedDocument:
    cleaned_segments = [segment for segment in segments if segment.text.strip()]
    text = "\n\n".join(segment.text for segment in cleaned_segments)
    return ExtractedDocument(text=text, segments=cleaned_segments)


def extract_document_from_bytes(
    content_bytes: bytes, extension: str
) -> ExtractedDocument:
    file_stream = io.BytesIO(content_bytes)
    try:
        normalized_extension = extension.lower()
        if normalized_extension == ".pdf":
            return extract_document_from_pdf(file_stream)
        if normalized_extension == ".docx":
            return extract_document_from_docx(file_stream)
        if normalized_extension in {".xls", ".xlsx"}:
            return extract_document_from_excel(file_stream, normalized_extension)
        if normalized_extension == ".txt":
            return extract_document_from_txt(file_stream)
        return ExtractedDocument(text="", segments=[])
    finally:
        file_stream.close()


def extract_document_from_path(path: str | Path) -> ExtractedDocument:
    file_path = Path(path)
    return extract_document_from_bytes(file_path.read_bytes(), file_path.suffix.lower())


def extract_document_from_pdf(file_obj: io.BytesIO) -> ExtractedDocument:
    document = _extract_pdf_with_pymupdf(file_obj)
    if document.text:
        return document

    return _extract_pdf_with_pypdf2(file_obj)


def _extract_pdf_with_pymupdf(file_obj: io.BytesIO) -> ExtractedDocument:
    try:
        import fitz

        pdf_bytes = file_obj.getvalue()
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        segments: list[ExtractedSegment] = []

        try:
            for page_index, page in enumerate(pdf_document, start=1):
                blocks = page.get_text("blocks", sort=True)
                block_texts = [
                    clean_extracted_text(str(block[4] or ""))
                    for block in blocks
                    if len(block) >= 5 and str(block[4] or "").strip()
                ]
                page_text = "\n\n".join(block for block in block_texts if block)
                if not page_text:
                    page_text = clean_extracted_text(page.get_text("text", sort=True))
                if not page_text:
                    continue
                segments.append(
                    ExtractedSegment(
                        text=page_text,
                        page_start=page_index,
                        page_end=page_index,
                    )
                )
        finally:
            pdf_document.close()

        return _build_document(segments)
    except Exception:
        return ExtractedDocument(text="", segments=[])


def _extract_pdf_with_pypdf2(file_obj: io.BytesIO) -> ExtractedDocument:
    try:
        import PyPDF2

        with contextlib.redirect_stderr(io.StringIO()):
            file_obj.seek(0)
            reader = PyPDF2.PdfReader(file_obj)
            segments: list[ExtractedSegment] = []
            for page_index, page in enumerate(reader.pages, start=1):
                page_text = clean_extracted_text(page.extract_text() or "")
                if not page_text:
                    continue
                segments.append(
                    ExtractedSegment(
                        text=page_text,
                        page_start=page_index,
                        page_end=page_index,
                    )
                )
        return _build_document(segments)
    except Exception:
        return ExtractedDocument(text="", segments=[])


def extract_document_from_docx(file_obj: io.BytesIO) -> ExtractedDocument:
    try:
        from docx import Document

        document = Document(file_obj)
        text = clean_extracted_text(
            "\n".join(paragraph.text for paragraph in document.paragraphs)
        )
        if not text:
            return ExtractedDocument(text="", segments=[])
        return _build_document([ExtractedSegment(text=text)])
    except Exception:
        return ExtractedDocument(text="", segments=[])


def _extract_xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    try:
        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    except KeyError:
        return []

    namespace = {"ns": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    shared_strings: list[str] = []

    for string_item in root.findall("ns:si", namespace):
        fragments = [
            node.text or "" for node in string_item.findall(".//ns:t", namespace)
        ]
        shared_strings.append("".join(fragments).strip())

    return shared_strings


def _extract_xlsx_sheet_names(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    workbook_path = "xl/workbook.xml"
    relations_path = "xl/_rels/workbook.xml.rels"
    namespace = {
        "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
        "rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "pkg": "http://schemas.openxmlformats.org/package/2006/relationships",
    }

    try:
        workbook_root = ET.fromstring(archive.read(workbook_path))
        relations_root = ET.fromstring(archive.read(relations_path))
    except KeyError:
        return []

    relation_targets: dict[str, str] = {}
    for relation in relations_root.findall("pkg:Relationship", namespace):
        relation_id = relation.attrib.get("Id")
        target = relation.attrib.get("Target")
        if not relation_id or not target:
            continue
        normalized_target = target.lstrip("/")
        if not normalized_target.startswith("xl/"):
            normalized_target = f"xl/{normalized_target}"
        relation_targets[relation_id] = normalized_target

    sheets: list[tuple[str, str]] = []
    for sheet in workbook_root.findall("main:sheets/main:sheet", namespace):
        sheet_name = (sheet.attrib.get("name") or "").strip() or "Sheet"
        relation_id = sheet.attrib.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        )
        sheet_path = relation_targets.get(relation_id or "")
        if sheet_path:
            sheets.append((sheet_name, sheet_path))

    return sheets


def _extract_xlsx_cell_value(
    cell: ET.Element,
    *,
    shared_strings: list[str],
    namespace: dict[str, str],
) -> str:
    cell_type = cell.attrib.get("t")
    raw_value = cell.findtext("ns:v", default="", namespaces=namespace).strip()

    if cell_type == "s":
        if raw_value.isdigit():
            shared_string_index = int(raw_value)
            if 0 <= shared_string_index < len(shared_strings):
                return shared_strings[shared_string_index]
        return ""

    if cell_type == "inlineStr":
        inline_text = "".join(
            node.text or "" for node in cell.findall(".//ns:is//ns:t", namespace)
        )
        return inline_text.strip()

    if cell_type == "b":
        return "TRUE" if raw_value == "1" else "FALSE"

    if cell_type == "str":
        return raw_value

    formula = cell.findtext("ns:f", default="", namespaces=namespace).strip()
    if formula and raw_value:
        return f"{formula} = {raw_value}"

    return raw_value


def _extract_xlsx_fallback(file_obj: io.BytesIO) -> str:
    namespace = {"ns": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}

    try:
        with zipfile.ZipFile(file_obj) as archive:
            shared_strings = _extract_xlsx_shared_strings(archive)
            sheets = _extract_xlsx_sheet_names(archive)
            if not sheets:
                sheets = [
                    (Path(name).stem, name)
                    for name in sorted(archive.namelist())
                    if name.startswith("xl/worksheets/") and name.endswith(".xml")
                ]

            rendered_sheets: list[str] = []
            for sheet_name, sheet_path in sheets:
                try:
                    sheet_root = ET.fromstring(archive.read(sheet_path))
                except KeyError:
                    continue

                rendered_rows: list[str] = []
                for row in sheet_root.findall(".//ns:sheetData/ns:row", namespace):
                    row_values = []
                    for cell in row.findall("ns:c", namespace):
                        value = _extract_xlsx_cell_value(
                            cell,
                            shared_strings=shared_strings,
                            namespace=namespace,
                        )
                        if value:
                            row_values.append(value)

                    if row_values:
                        rendered_rows.append(" | ".join(row_values))

                if rendered_rows:
                    rendered_sheets.append(
                        f"Sheet: {sheet_name}\n" + "\n".join(rendered_rows)
                    )

            return "\n\n".join(rendered_sheets)
    except zipfile.BadZipFile:
        return ""
    except Exception:
        return ""


def extract_document_from_excel(
    file_obj: io.BytesIO,
    extension: str,
) -> ExtractedDocument:
    try:
        dataframe_by_sheet = pd.read_excel(file_obj, sheet_name=None)
    except Exception:
        dataframe_by_sheet = None

    if dataframe_by_sheet is not None:
        output: list[str] = []
        for sheet_name, sheet in dataframe_by_sheet.items():
            cleaned_sheet = sheet.dropna(how="all").dropna(axis=1, how="all")
            if cleaned_sheet.empty:
                continue

            rendered_sheet = (
                cleaned_sheet.fillna("")
                .astype(str)
                .apply(
                    lambda row: " | ".join(
                        value.strip()
                        for value in row
                        if isinstance(value, str) and value.strip()
                    ),
                    axis=1,
                )
            )
            rendered_rows = [row for row in rendered_sheet.tolist() if row]
            if rendered_rows:
                output.append(f"Sheet: {sheet_name}\n" + "\n".join(rendered_rows))

        text = clean_extracted_text("\n\n".join(output))
        if text:
            return _build_document([ExtractedSegment(text=text)])

    if extension == ".xlsx":
        file_obj.seek(0)
        text = clean_extracted_text(_extract_xlsx_fallback(file_obj))
        if text:
            return _build_document([ExtractedSegment(text=text)])

    return ExtractedDocument(text="", segments=[])


def extract_document_from_txt(file_obj: io.BytesIO) -> ExtractedDocument:
    try:
        text = clean_extracted_text(file_obj.read().decode("utf-8", errors="ignore"))
        if not text:
            return ExtractedDocument(text="", segments=[])
        return _build_document([ExtractedSegment(text=text)])
    except Exception:
        return ExtractedDocument(text="", segments=[])


def split_long_piece(piece: str, *, max_chars: int) -> list[str]:
    if len(piece) <= max_chars:
        return [piece]

    chunks: list[str] = []
    start = 0
    piece_length = len(piece)
    while start < piece_length:
        end = min(start + max_chars, piece_length)
        chunks.append(piece[start:end])
        start = end
    return chunks


def iter_text_blocks(text: str, *, max_block_chars: int) -> list[str]:
    if not text:
        return []

    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    raw_blocks = [block.strip() for block in normalized.split("\n\n") if block.strip()]
    if not raw_blocks:
        raw_blocks = [normalized.strip()]

    blocks: list[str] = []
    for block in raw_blocks:
        if len(block) <= max_block_chars:
            blocks.append(block)
            continue

        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if len(lines) > 1:
            for line in lines:
                blocks.extend(split_long_piece(line, max_chars=max_block_chars))
        else:
            blocks.extend(split_long_piece(block, max_chars=max_block_chars))

    return [block for block in blocks if block]


class StreamingChunker:
    def __init__(self, tokenizer, *, chunk_tokens: int, chunk_overlap: int):
        self.tokenizer = tokenizer
        self.chunk_tokens = chunk_tokens
        self.chunk_overlap = chunk_overlap
        self.step_tokens = max(chunk_tokens - chunk_overlap, 1)

    def chunk_document(
        self,
        document: ExtractedDocument,
        *,
        max_block_chars: int,
    ) -> list[ChunkedSegment]:
        token_buffer: list[int] = []
        page_buffer: list[tuple[int | None, int | None]] = []
        output_chunks: list[ChunkedSegment] = []

        for segment in document.segments:
            for block in iter_text_blocks(
                segment.text, max_block_chars=max_block_chars
            ):
                block_tokens = self.tokenizer.encode(block, add_special_tokens=False)
                if not block_tokens:
                    continue

                token_buffer.extend(block_tokens)
                page_buffer.extend(
                    [(segment.page_start, segment.page_end)] * len(block_tokens)
                )

                while len(token_buffer) >= self.chunk_tokens:
                    output_chunks.append(
                        self._build_chunk(
                            chunk_index=len(output_chunks),
                            token_buffer=token_buffer[: self.chunk_tokens],
                            page_buffer=page_buffer[: self.chunk_tokens],
                        )
                    )
                    token_buffer = token_buffer[self.step_tokens :]
                    page_buffer = page_buffer[self.step_tokens :]

        if token_buffer:
            output_chunks.append(
                self._build_chunk(
                    chunk_index=len(output_chunks),
                    token_buffer=token_buffer,
                    page_buffer=page_buffer,
                )
            )

        return [chunk for chunk in output_chunks if chunk.text]

    def chunk_text(self, text: str, *, max_block_chars: int) -> list[str]:
        document = ExtractedDocument(
            text=text,
            segments=[ExtractedSegment(text=text)] if text else [],
        )
        return [
            chunk.text
            for chunk in self.chunk_document(document, max_block_chars=max_block_chars)
        ]

    def _build_chunk(
        self,
        *,
        chunk_index: int,
        token_buffer: list[int],
        page_buffer: list[tuple[int | None, int | None]],
    ) -> ChunkedSegment:
        decoded = self.tokenizer.decode(token_buffer).strip()
        page_start = self._first_page(page_buffer)
        page_end = self._last_page(page_buffer)
        return ChunkedSegment(
            chunk_index=chunk_index,
            text=decoded,
            page_start=page_start,
            page_end=page_end,
            source_locator=format_source_locator(page_start, page_end),
        )

    @staticmethod
    def _first_page(page_buffer: list[tuple[int | None, int | None]]) -> int | None:
        for page_start, _ in page_buffer:
            if page_start:
                return page_start
        return None

    @staticmethod
    def _last_page(page_buffer: list[tuple[int | None, int | None]]) -> int | None:
        for _, page_end in reversed(page_buffer):
            if page_end:
                return page_end
        return None
